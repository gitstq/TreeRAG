"""
文档解析器模块

支持多种文档格式的解析，包括PDF、Markdown、TXT和DOCX。
将文档内容智能分割为适合索引的段落片段。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from .models import Document

logger = logging.getLogger(__name__)


class DocumentParser:
    """多格式文档解析器。

    自动检测文件格式并解析为统一的Document对象列表。
    支持PDF、Markdown、TXT、DOCX四种格式。

    Usage:
        >>> parser = DocumentParser(max_segment_length=1000)
        >>> documents = parser.parse("example.pdf")
        >>> for doc in documents:
        ...     print(doc.content[:100])
    """

    # 支持的文件扩展名映射
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "txt",
        ".text": "txt",
        ".docx": "docx",
    }

    def __init__(self, max_segment_length: int = 1000) -> None:
        """初始化文档解析器。

        Args:
            max_segment_length: 分段最大字符长度
        """
        self.max_segment_length = max_segment_length

    def parse(self, file_path: str) -> list[Document]:
        """解析文件，自动检测格式。

        Args:
            file_path: 文件路径

        Returns:
            解析后的Document对象列表

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        file_type = self.SUPPORTED_EXTENSIONS.get(suffix)

        if file_type is None:
            raise ValueError(
                f"不支持的文件格式: {suffix}。"
                f"支持的格式: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
            )

        logger.info(f"正在解析文件: {file_path} (格式: {file_type})")

        if file_type == "pdf":
            documents = self.parse_pdf(str(path))
        elif file_type == "markdown":
            documents = self.parse_markdown(str(path))
        elif file_type == "txt":
            documents = self.parse_txt(str(path))
        elif file_type == "docx":
            documents = self.parse_docx(str(path))
        else:
            documents = []

        logger.info(f"解析完成，共获得 {len(documents)} 个文档片段")
        return documents

    def parse_pdf(self, path: str) -> list[Document]:
        """解析PDF文件。

        使用PyPDF2提取文本内容，按页分割后进一步按段落分段。

        Args:
            path: PDF文件路径

        Returns:
            解析后的Document对象列表
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError(
                "解析PDF需要安装PyPDF2库。请运行: pip install PyPDF2"
            )

        documents: list[Document] = []
        reader = PdfReader(path)

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if not text or not text.strip():
                continue

            # 清理PDF提取的文本
            text = self._clean_text(text)

            # 分段
            segments = self._split_into_segments(text)
            for seg in segments:
                if seg.strip():
                    documents.append(Document(
                        content=seg.strip(),
                        metadata={
                            "source": path,
                            "page": page_num + 1,
                            "format": "pdf",
                        },
                    ))

        return documents

    def parse_markdown(self, path: str) -> list[Document]:
        """解析Markdown文件。

        按标题层级分割内容，保留层级信息作为元数据。

        Args:
            path: Markdown文件路径

        Returns:
            解析后的Document对象列表
        """
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        documents: list[Document] = []
        # 按标题分割
        sections = re.split(r'^(#{1,6}\s+.+)$', text, flags=re.MULTILINE)

        current_section = "引言"
        for i, section in enumerate(sections):
            section = section.strip()
            if not section:
                continue

            # 检查是否是标题行
            if re.match(r'^#{1,6}\s+', section):
                current_section = re.sub(r'^#{1,6}\s+', '', section).strip()
                continue

            # 分段处理
            segments = self._split_into_segments(section)
            for seg in segments:
                if seg.strip():
                    documents.append(Document(
                        content=seg.strip(),
                        metadata={
                            "source": path,
                            "section": current_section,
                            "format": "markdown",
                        },
                    ))

        # 如果没有按标题分割出内容，则整体处理
        if not documents and text.strip():
            segments = self._split_into_segments(text)
            for seg in segments:
                if seg.strip():
                    documents.append(Document(
                        content=seg.strip(),
                        metadata={
                            "source": path,
                            "section": "全文",
                            "format": "markdown",
                        },
                    ))

        return documents

    def parse_txt(self, path: str) -> list[Document]:
        """解析纯文本文件。

        按段落和换行符分割内容。

        Args:
            path: 文本文件路径

        Returns:
            解析后的Document对象列表
        """
        # 尝试多种编码
        text: Optional[str] = None
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(path, "r", encoding=encoding) as f:
                    text = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if text is None:
            raise ValueError(f"无法解码文件: {path}，尝试了utf-8、gbk、gb2312、latin-1编码")

        text = self._clean_text(text)
        segments = self._split_into_segments(text)

        documents: list[Document] = []
        for seg in segments:
            if seg.strip():
                documents.append(Document(
                    content=seg.strip(),
                    metadata={
                        "source": path,
                        "format": "txt",
                    },
                ))

        return documents

    def parse_docx(self, path: str) -> list[Document]:
        """解析DOCX文件。

        使用python-docx提取文本内容，按段落分割。

        Args:
            path: DOCX文件路径

        Returns:
            解析后的Document对象列表
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "解析DOCX需要安装python-docx库。请运行: pip install python-docx"
            )

        doc = DocxDocument(path)
        full_text: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        combined_text = "\n\n".join(full_text)
        segments = self._split_into_segments(combined_text)

        documents: list[Document] = []
        for seg in segments:
            if seg.strip():
                documents.append(Document(
                    content=seg.strip(),
                    metadata={
                        "source": path,
                        "format": "docx",
                    },
                ))

        return documents

    def _split_into_segments(self, text: str) -> list[str]:
        """智能分段，尊重句子边界。

        将长文本分割为不超过max_segment_length的段落，
        尽量在句子边界处分割。

        Args:
            text: 要分割的文本

        Returns:
            分割后的文本段落列表
        """
        if not text or not text.strip():
            return []

        text = text.strip()

        # 如果文本较短，直接返回
        if len(text) <= self.max_segment_length:
            return [text]

        segments: list[str] = []

        # 先按段落分割
        paragraphs = re.split(r'\n\s*\n', text)
        current_segment = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # 如果当前段落本身超过最大长度，需要进一步分割
            if len(para) > self.max_segment_length:
                # 先保存当前累积的段落
                if current_segment:
                    segments.append(current_segment)
                    current_segment = ""

                # 按句子分割长段落
                sentences = re.split(r'(?<=[。！？.!?\n])', para)
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue

                    if len(current_segment) + len(sentence) + 1 <= self.max_segment_length:
                        current_segment = (
                            current_segment + " " + sentence
                            if current_segment
                            else sentence
                        )
                    else:
                        if current_segment:
                            segments.append(current_segment)
                        current_segment = sentence
            else:
                # 段落不超过最大长度
                if len(current_segment) + len(para) + 2 <= self.max_segment_length:
                    current_segment = (
                        current_segment + "\n\n" + para
                        if current_segment
                        else para
                    )
                else:
                    if current_segment:
                        segments.append(current_segment)
                    current_segment = para

        # 添加最后一个段落
        if current_segment:
            segments.append(current_segment)

        return segments

    @staticmethod
    def _clean_text(text: str) -> str:
        """清理文本内容。

        去除多余的空白字符、特殊字符等。

        Args:
            text: 原始文本

        Returns:
            清理后的文本
        """
        # 替换各种空白字符为普通空格
        text = re.sub(r'[\t\r\f\v]', ' ', text)
        # 合并多个连续空格
        text = re.sub(r' {3,}', '  ', text)
        # 合并多个连续空行为两个换行
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 去除行首行尾空白
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        return text.strip()
